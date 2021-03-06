# Testmodule
import docx
from docx.shared import Inches


class Text:

    def __init__(self, string, bold=False, italic=False):
        self.string = string
        self.bold = bold
        self.italic = italic


class NewDocument:
    """Get directions between locations

        :param origin: Origin location - string address; (latitude, longitude)
            two-tuple, dict with ("lat", "lon") keys or object with (lat, lon)
            attributes
        :param destination: Destination location - type same as origin
        :param mode: Travel mode as string, defaults to "driving".
            See `google docs details <https://developers.google.com/maps/documentation/directions/#TravelModes>`_
        :param alternatives: True if provide it has to return more then one
            route alternative
        :param waypoints: Iterable with set of intermediate stops,
            like ("Munich", "Dallas")
            See `google docs details <https://developers.google.com/maps/documentation/javascript/reference#DirectionsRequest>`_
        :param optimize_waypoints: if true will attempt to re-order supplied
            waypoints to minimize overall cost of the route. If waypoints are
            optimized, the route returned will show the optimized order under
            "waypoint_order". See `google docs details <https://developers.google.com/maps/documentation/javascript/reference#DirectionsRequest>`_
        :param avoid: Iterable with set of restrictions,
            like ("tolls", "highways"). For full list refer to
            `google docs details <https://developers.google.com/maps/documentation/directions/#Restrictions>`_
        :param language: The language in which to return results.
            See `list of supported languages <https://developers.google.com/maps/faq#languagesupport>`_
        :param units: Unit system for result. Defaults to unit system of
            origin's country.
            See `google docs details <https://developers.google.com/maps/documentation/directions/#UnitSystems>`_
        :param region: The region code. Affects geocoding of origin and
            destination (see `gmaps.Geocoding.geocode` region parameter)
        :param departure_time: Desired time of departure as
            seconds since midnight, January 1, 1970 UTC
        :param arrival_time: Desired time of arrival for transit directions as
            seconds since midnight, January 1, 1970 UTC.
        """

    def __init__(self, name="mydocx"):
        self.doc = docx.Document()
        self.name = name
        self.docdata = []

    def __str__(self):
        amount = str(len(self.docdata))
        return "docdata now conatains {items} items".format(items=amount)

    def getInsertPoint(self):
        """
        Returns current length to be used as InsertPoint
        """
        return len(self.docdata)

    def save(self):
        """
        Loops through all items in self.docdata and adds them to the self.doc (python-docx) object.
        Then saves document with name self.name
        """

        def save_text(self, item):

            string = item["textdata"].string  # string
            bold = item["textdata"].bold  # boolean
            italic = item["textdata"].italic  # boolean

            if bold:
                p = self.doc.add_paragraph('')
                p.add_run(string).bold = True
            elif italic:
                p = self.doc.add_paragraph('')
                p.add_run(string).italic = True
            else:
                self.doc.add_paragraph(string)

        def save_multitext(self, item):
            strings = [i.string for i in item["textdata"]]  # array string
            bolds = [i.bold for i in item["textdata"]]  # array boolean
            italics = [i.italic for i in item["textdata"]]  # array boolean

            for ind, string in enumerate(strings):

                if ind == 0:
                    if bolds[ind]:
                        p = self.doc.add_paragraph('')
                        p.add_run(string).bold = True
                    elif italics[ind]:
                        p = self.doc.add_paragraph('')
                        p.add_run(string).italic = True
                    else:
                        p = self.doc.add_paragraph(string)
                else:
                    if bolds[ind]:
                        p.add_run(string).bold = True
                    elif italics[ind]:
                        p.add_run(string).italic = True
                    else:
                        p.add_run(string)

        def save_heading(self, item):
            string = item["string"]
            level = item["level"]

            self.doc.add_heading(string, level)

        def save_picture(self, item):
            path = item["path"]
            width = item["width"]

            self.doc.add_picture(path, width)

        for item in self.docdata:
            if item["category"] == "text":
                save_text(self, item)
            if item["category"] == "multitext":
                save_multitext(self, item)
            elif item["category"] == "heading":
                save_heading(self, item)
            elif item["category"] == "picture":
                save_picture(self, item)

        self.doc.save(self.name+".docx")

    def addText(self, textobj, insertPoint=False):
        """
        Stores a dictionary of category text to self.docdata.
        """
        if not insertPoint:
            insertPoint = len(self.docdata)

        self.docdata.insert(insertPoint, {
            "category": "text",
            "textdata": textobj
        })

    def addMultiText(self, textobjs, insertPoint=False):
        """
        Stores a dictionary of category multitext to self.docdata.
        """
        if not insertPoint:
            insertPoint = len(self.docdata)

        self.docdata.insert(insertPoint, {
            "category": "multitext",
            "textdata": textobjs
        })

    def addHeading(self, string, level=1, insertPoint=False):
        """
        Stores a dictionary of category heading to self.docdata.
        """
        self.docdata.append({
            "category": "heading",
            "string": string,
            "level": level
        })

    def addPicture(self, path, width=Inches(6.3), insertPoint=False):
        """
        Stores a dictionary of category picture to self.docdata.
        """
        self.docdata.append({
            "category": "picture",
            "path": path,
            "width": width
        })

    def test(self):
        document = docx.Document()


if __name__ == '__main__':
    import sys
    sys.path.append(r"H:\Python\mypythonmodules")

    from avbr_docx import NewDocument, Text

    doc = NewDocument("hello")

    insertPoints = {}

    doc.addText(Text("Hello", bold=True, italic=True))

    insertPoints["point1"] = doc.getInsertPoint()

    doc.addMultiText([Text("Hello ", bold=True, italic=True),
                      Text("World", bold=False, italic=True)])

    # doc.addPicture("test.jpg")

    doc.addText(Text("Hello", bold=True, italic=True),
                insertPoint=insertPoints["point1"])

    doc.addHeading("Hello", level=2)

    doc.save()
